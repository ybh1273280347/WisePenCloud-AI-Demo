from __future__ import annotations

import sys
import tempfile
import types
import unittest
from pathlib import Path

sys.path.insert(0, "src")

from chat.application.document_parse.office.native_parser import OfficeNativeParser
from chat.application.document_parse.pdf.table_extractor import TableExtractor
from chat.application.document_parse.spreadsheet.parser import SpreadsheetParser


class DocumentParseQualityTest(unittest.TestCase):
    def test_spreadsheet_rows_preserve_na_text_and_clean_tsv_cells(self) -> None:
        import pandas as pd

        df = pd.DataFrame(
            [["NA", "N/A", "NULL", "line 1\nline 2", "a\tb", "  many   spaces  "]],
            columns=["code", "short", "null_text", "notes", "tabbed", "spaced"],
        )

        rows = SpreadsheetParser()._rows_from_dataframe(df)

        self.assertEqual(rows[1][0:3], ["NA", "N/A", "NULL"])
        self.assertEqual(rows[1][3], "line 1 / line 2")
        self.assertEqual(rows[1][4], "a b")
        self.assertEqual(rows[1][5], "many spaces")
        self.assertNotIn("\n", "\t".join(rows[1]))

    def test_spreadsheet_none_cell_becomes_empty_string(self) -> None:
        import pandas as pd

        df = pd.DataFrame(
            [[None, "", "text"]],
            columns=["none_col", "empty_col", "text_col"],
        )

        rows = SpreadsheetParser()._rows_from_dataframe(df)

        self.assertEqual(rows[1][0], "", "None cell must become empty string, not 'None'")
        self.assertEqual(rows[1][1], "", "empty cell must remain empty string")
        self.assertEqual(rows[1][2], "text")

    def test_camelot_stream_only_runs_when_lattice_has_no_valid_rows(self) -> None:
        import pandas as pd

        calls = []

        class FakeTable:
            def __init__(self, values):
                self.df = pd.DataFrame(values)

        fake_camelot = types.SimpleNamespace()

        def read_pdf(_path: str, *, pages: str, flavor: str):
            calls.append((pages, flavor))
            if flavor == "lattice":
                return [FakeTable([["", ""]])]
            return [FakeTable([["A", "B"]])]

        fake_camelot.read_pdf = read_pdf
        previous = sys.modules.get("camelot")
        sys.modules["camelot"] = fake_camelot
        try:
            tables = TableExtractor().extract_tables(Path("sample.pdf"), page_index=2)
        finally:
            if previous is None:
                sys.modules.pop("camelot", None)
            else:
                sys.modules["camelot"] = previous

        self.assertEqual(calls, [("3", "lattice"), ("3", "stream")])
        self.assertEqual(tables[0].rows, [["A", "B"]])
        self.assertEqual(tables[0].metadata["flavor"], "stream")

    def test_docx_native_fallback_preserves_paragraph_table_order(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "ordered.docx"
            doc = Document()
            doc.add_paragraph("before table")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "inside table"
            doc.add_paragraph("after table")
            doc.save(path)

            result = OfficeNativeParser().parse(path, file_type="docx")

        before_index = result.text.index("before table")
        table_index = result.text.index("[Table 1]")
        after_index = result.text.index("after table")

        self.assertLess(before_index, table_index)
        self.assertLess(table_index, after_index)
        self.assertEqual(result.tables[0].rows, [["inside table"]])


if __name__ == "__main__":
    unittest.main()
